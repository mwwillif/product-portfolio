from docx import Document
import re

SECTION_MARKERS = {
    "summary": ("[[SUMMARY_START]]", "[[SUMMARY_END]]"),
    "skills_data_platforms": ("[[SKILLS_DATA_PLATFORMS_START]]", "[[SKILLS_DATA_PLATFORMS_END]]"),
    "skills_product_leadership": ("[[SKILLS_PRODUCT_LEADERSHIP_START]]", "[[SKILLS_PRODUCT_LEADERSHIP_END]]"),
    "skills_tools": ("[[SKILLS_TOOLS_START]]", "[[SKILLS_TOOLS_END]]"),
    "skills_domains": ("[[SKILLS_DOMAINS_START]]", "[[SKILLS_DOMAINS_END]]"),
    "amplytico_desc": ("[[AMPLYTICO_DESC_START]]", "[[AMPLYTICO_DESC_END]]"),
    "amplytico_bullets": ("[[AMPLYTICO_BULLETS_START]]", "[[AMPLYTICO_BULLETS_END]]"),
    "wells_fargo_vp_desc": ("[[WELLS_FARGO_VP_DESC_START]]", "[[WELLS_FARGO_VP_DESC_END]]"),
    "wells_fargo_vp_bullets": ("[[WELLS_FARGO_VP_BULLETS_START]]", "[[WELLS_FARGO_VP_BULLETS_END]]"),
    "ey_desc": ("[[EY_DESC_START]]", "[[EY_DESC_END]]"),
    "ey_bullets": ("[[EY_BULLETS_START]]", "[[EY_BULLETS_END]]"),
    "wells_fargo_reg_desc": ("[[WELLS_FARGO_REG_DESC_START]]", "[[WELLS_FARGO_REG_DESC_END]]"),
    "wells_fargo_reg_bullets": ("[[WELLS_FARGO_REG_BULLETS_START]]", "[[WELLS_FARGO_REG_BULLETS_END]]"),
}

ALL_MARKERS = set()
for start_marker, end_marker in SECTION_MARKERS.values():
    ALL_MARKERS.add(start_marker)
    ALL_MARKERS.add(end_marker)

MARKER_REGEX = re.compile(r"\[\[[A-Z0-9_]+(?:_START|_END)\]\]")

SKILL_LABELS = {
    "skills_data_platforms": "Data Platforms & Governance:",
    "skills_product_leadership": "Product Leadership:",
    "skills_tools": "Tools:",
    "skills_domains": "Domains:",
}


def clear_paragraph(paragraph):
    for run in paragraph.runs:
        run.text = ""
    if not paragraph.runs:
        paragraph.text = ""


def write_block_paragraph_text(paragraph, text: str):
    """
    Only used for block sections inside markers.
    This does not touch header/date lines because those are outside markers.
    """
    clear_paragraph(paragraph)
    if paragraph.runs:
        paragraph.runs[0].text = text
    else:
        paragraph.add_run(text)


def delete_paragraph(paragraph):
    p = paragraph._element
    parent = p.getparent()
    if parent is not None:
        parent.remove(p)


def set_inline_skill_paragraph(paragraph, label, value):
    """
    Rebuild one skills paragraph as:
    Bold label + normal value on same line
    """
    clear_paragraph(paragraph)

    label_run = paragraph.add_run(label)
    label_run.bold = True

    value_run = paragraph.add_run(" " + value)
    value_run.bold = False


def replace_inline_skill_paragraph(paragraph, section_id, value):
    """
    Handles only skill lines like:
    Data Platforms & Governance: [[SKILLS_DATA_PLATFORMS_START]]...[[SKILLS_DATA_PLATFORMS_END]]
    """
    start_marker, end_marker = SECTION_MARKERS[section_id]
    text = paragraph.text

    if start_marker not in text or end_marker not in text:
        return False

    set_inline_skill_paragraph(paragraph, SKILL_LABELS[section_id], value)
    return True


def find_paragraph_index(paragraphs, marker):
    for i, p in enumerate(paragraphs):
        if p.text.strip() == marker:
            return i
    return None


def apply_updates(template, output, updates):
    doc = Document(template)
    updates_map = {u["section_id"]: u["content"] for u in updates["updates"]}

    # 1) Skills only: replace inline marker paragraphs
    for section_id in ("skills_data_platforms", "skills_product_leadership", "skills_tools", "skills_domains"):
        if section_id not in updates_map:
            continue

        value = updates_map[section_id][0] if updates_map[section_id] else ""

        for p in doc.paragraphs:
            if replace_inline_skill_paragraph(p, section_id, value):
                break

    # 2) Block sections only: desc + bullets between markers
    paragraphs = doc.paragraphs

    for section_id, (start_marker, end_marker) in SECTION_MARKERS.items():
        if section_id.startswith("skills_"):
            continue
        if section_id not in updates_map:
            continue

        start_idx = find_paragraph_index(paragraphs, start_marker)
        end_idx = find_paragraph_index(paragraphs, end_marker)

        if start_idx is None or end_idx is None or end_idx <= start_idx:
            continue

        existing = paragraphs[start_idx + 1:end_idx]
        new_lines = updates_map[section_id]

        # Reuse only existing paragraphs. Do not insert.
        for i, p in enumerate(existing):
            if i < len(new_lines):
                write_block_paragraph_text(p, new_lines[i])
            else:
                write_block_paragraph_text(p, "")

    # 3) Remove marker-only paragraphs
    for p in list(doc.paragraphs):
        if p.text.strip() in ALL_MARKERS:
            delete_paragraph(p)

    doc.save(output)
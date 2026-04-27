from docx import Document
import os

SECTION_MARKERS = {
    "summary": ("[[SUMMARY_START]]", "[[SUMMARY_END]]"),
    "skills_data_platforms": ("[[SKILLS_DATA_PLATFORMS_START]]", "[[SKILLS_DATA_PLATFORMS_END]]"),
    "skills_product_leadership": ("[[SKILLS_PRODUCT_LEADERSHIP_START]]", "[[SKILLS_PRODUCT_LEADERSHIP_END]]"),
    "skills_tools": ("[[SKILLS_TOOLS_START]]", "[[SKILLS_TOOLS_END]]"),
    "skills_domains": ("[[SKILLS_DOMAINS_START]]", "[[SKILLS_DOMAINS_END]]"),
    "amplytico_desc": ("[[AMPLYTICO_DESC_START]]", "[[AMPLYTICO_DESC_END]]"),
    "amplytico_bullets": ("[[AMPLYTICO_BULLETS_START]]", "[[AMPLYTICO_BULLETS_END]]"),
    "wells_fargo_vp_desc": ("[[WELLS_FARGO_VP_DESC_START]]", "[[WELLS_FARGO_VP_DESC_END]]"),
    "wells_fargo_vp_bullets": ("[[WELLS_FARGO_VP_BULLETS_START]]", "[[WELLS_FARGO_VP_BULLETS_END]]"),
    "ey_desc": ("[[EY_DESC_START]]", "[[EY_DESC_END]]"),
    "ey_bullets": ("[[EY_BULLETS_START]]", "[[EY_BULLETS_END]]"),
    "wells_fargo_reg_desc": ("[[WELLS_FARGO_REG_DESC_START]]", "[[WELLS_FARGO_REG_DESC_END]]"),
    "wells_fargo_reg_bullets": ("[[WELLS_FARGO_REG_BULLETS_START]]", "[[WELLS_FARGO_REG_BULLETS_END]]"),
}

ALL_MARKERS = set()
for s, e in SECTION_MARKERS.values():
    ALL_MARKERS.add(s)
    ALL_MARKERS.add(e)

LABEL_MARKERS = {
    "[[SKILLS_DATA_PLATFORMS_LABEL]]",
    "[[SKILLS_PRODUCT_LEADERSHIP_LABEL]]",
    "[[SKILLS_TOOLS_LABEL]]",
    "[[SKILLS_DOMAINS_LABEL]]",
}

def extract_template_sections(docx_path):
    doc = Document(docx_path)
    sections = {}
    current = None
    buffer = []

    start_lookup = {v[0]: k for k, v in SECTION_MARKERS.items()}
    end_lookup = {v[1]: k for k, v in SECTION_MARKERS.items()}

    for p in doc.paragraphs:
        text = p.text.strip()

        if text in LABEL_MARKERS:
            continue

        if text in start_lookup:
            current = start_lookup[text]
            buffer = []
            continue

        if text in end_lookup:
            sections[current] = [line for line in buffer if line]
            current = None
            buffer = []
            continue

        if current:
            buffer.append(text)

    return sections

def extract_full_text(docx_path):
    doc = Document(docx_path)
    lines = []

    for p in doc.paragraphs:
        t = p.text.strip()
        if not t:
            continue
        if t in ALL_MARKERS or t in LABEL_MARKERS:
            continue
        lines.append(t)

    return "\n".join(lines)

def load_all_experience(input_folder, template_name):
    combined = []

    for f in os.listdir(input_folder):
        if not f.endswith(".docx"):
            continue
        if f == template_name:
            continue

        path = os.path.join(input_folder, f)
        combined.append(extract_full_text(path))

    return "\n\n".join(combined)